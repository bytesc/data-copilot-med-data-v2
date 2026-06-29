import time
from typing import Optional, List
import io
import base64
import httpx
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pywebio.session import set_env
from pywebio.input import input, TEXT, textarea, file_upload, select, checkbox
from pywebio.output import put_text, put_html, put_markdown, clear, put_loading, toast, popup, put_buttons, \
    put_collapse, put_table, put_scope, use_scope, put_row
from pywebio import start_server, config
from data_access.read_db import get_rows_from_all_tables, get_table_comments_dict, get_all_comments_from_table
from utils.get_config import config_data
import markdown
from bs4 import BeautifulSoup
import requests
from urllib.parse import urlparse

SELECT_TABLES = []
SELECT_LABELS = []


def ai_agent_api(question: str, tables: Optional[List[str]] = None, path: str = "/api/ask-agent/",
                 url="http://127.0.0.1:" + str(config_data["server_port"])):
    with httpx.Client(timeout=180.0) as client:
        try:
            payload = {"question": question}
            if tables:
                payload["tables"] = tables

            response = client.post(url + path, json=payload)
            if response.status_code == 200:
                print(response.json()["ans"])
                return response.json()["ans"], response.json()["code"]
            else:
                return None
        except httpx.RequestError as e:
            print(e)
            return None


def upload_csv_api(file_content, table_name="uploaded_data"):
    url = f"http://127.0.0.1:{config_data['server_port']}/upload-csv/"
    files = {
        'file': ('data.csv', file_content, 'text/csv')
    }
    data = {
        'table_name': table_name
    }
    with httpx.Client(timeout=30.0) as client:
        try:
            response = client.post(url, files=files, data=data)
            if response.status_code == 200:
                return response.json()
            else:
                return {"error": f"{response.status_code}", "details": response.text}
        except httpx.RequestError as e:
            return {"error": f"{str(e)}"}


def upload_doc_api(file_content, filename, table_name="uploaded_data"):
    url = f"http://127.0.0.1:{config_data['server_port']}/upload-txt/"
    files = {
        'file': (filename, file_content, 'application/octet-stream')
    }
    data = {
        'table_name': table_name
    }
    with httpx.Client(timeout=30.0) as client:
        try:
            response = client.post(url, files=files, data=data)
            if response.status_code == 200:
                return response.json()
            else:
                return {"error": f"{response.status_code}", "details": response.text}
        except httpx.RequestError as e:
            return {"error": f"{str(e)}"}


def download_image(url):
    """下载网络图片"""
    try:
        response = requests.get(url, timeout=10)
        if response.status_code == 200:
            return io.BytesIO(response.content)
    except:
        pass
    return None


def markdown_to_word(doc, markdown_text):
    """将Markdown转换为Word文档内容"""
    html = markdown.markdown(markdown_text, extensions=['extra', 'tables'])
    soup = BeautifulSoup(html, 'html.parser')

    for element in soup.children:
        if element.name == 'h1':
            doc.add_heading(element.get_text(), level=1)
        elif element.name == 'h2':
            doc.add_heading(element.get_text(), level=2)
        elif element.name == 'h3':
            doc.add_heading(element.get_text(), level=3)
        elif element.name == 'h4':
            doc.add_heading(element.get_text(), level=4)
        elif element.name == 'p':
            p = doc.add_paragraph()
            # 处理段落中的内联元素
            for child in element.children:
                if child.name == 'strong' or child.name == 'b':
                    run = p.add_run(child.get_text())
                    run.bold = True
                elif child.name == 'em' or child.name == 'i':
                    run = p.add_run(child.get_text())
                    run.italic = True
                elif child.name == 'a':
                    run = p.add_run(child.get_text())
                    run.underline = True
                    # 可以添加超链接
                    # doc.add_hyperlink(child.get_text(), child.get('href'))
                elif child.name == 'img':
                    img_url = child.get('src')
                    if img_url:
                        img_data = download_image(img_url)
                        if img_data:
                            try:
                                doc.add_picture(img_data, width=Inches(5))
                            except:
                                p.add_run(f"[图片: {img_url}]")
                        else:
                            p.add_run(f"[图片加载失败: {img_url}]")
                elif child.string:
                    p.add_run(child.string)
                elif child.name is None:
                    if child.strip():
                        p.add_run(child.strip())
        elif element.name == 'ul':
            for li in element.find_all('li', recursive=False):
                doc.add_paragraph(li.get_text(), style='List Bullet')
        elif element.name == 'ol':
            for idx, li in enumerate(element.find_all('li', recursive=False), 1):
                doc.add_paragraph(f"{idx}. {li.get_text()}", style='List Number')
        elif element.name == 'table':
            rows = element.find_all('tr')
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0].find_all(['td', 'th'])))
                table.style = 'Table Grid'
                for i, row in enumerate(rows):
                    cells = row.find_all(['td', 'th'])
                    for j, cell in enumerate(cells):
                        table.cell(i, j).text = cell.get_text().strip()
        elif element.name == 'blockquote':
            p = doc.add_paragraph()
            p.add_run(element.get_text()).italic = True
            p.paragraph_format.left_indent = Inches(0.5)
        elif element.name == 'hr':
            doc.add_page_break()
        elif element.name == 'br':
            doc.add_paragraph()
        elif element.string and element.string.strip():
            # 处理纯文本内容
            doc.add_paragraph(element.string.strip())


def export_full_to_word(conversation_history):
    """导出完整的对话历史为Word文档"""
    doc = Document()

    # 添加标题
    title = doc.add_heading('Data-Copilot Conversation Export (Full)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 添加导出时间
    from datetime import datetime
    doc.add_paragraph(f"Export Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph()

    # 处理每条对话
    for entry in conversation_history:
        if entry.startswith('Q: '):
            doc.add_heading('Question:', level=2)
            doc.add_paragraph(entry[3:])
        elif entry.startswith('A: '):
            doc.add_heading('Answer:', level=2)
            markdown_to_word(doc, entry[3:])
        elif entry.startswith('Code Generated: '):
            doc.add_heading('Generated Code:', level=2)
            code_para = doc.add_paragraph()
            code_run = code_para.add_run(entry[16:])
            code_run.font.name = 'Courier New'
            code_run.font.size = Pt(10)
        elif entry.startswith('Exe Result: '):
            doc.add_heading('Execution Result:', level=2)
            markdown_to_word(doc, entry[12:])
        elif entry.startswith('Planner: '):
            doc.add_heading('Plan:', level=2)
            markdown_to_word(doc, entry[9:])
        else:
            doc.add_paragraph(entry)

        # 添加分隔线
        doc.add_paragraph('_' * 50)

    # 保存到内存
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    return file_stream


def export_essentials_to_word(conversation_history):
    """只导出答案(ans)和第一个问题（去除标签，只输出内容，不包含代码）"""
    doc = Document()

    # 添加标题
    title = doc.add_heading('Data-Copilot Export', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 添加导出时间
    from datetime import datetime
    doc.add_paragraph(f"Export Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph()

    # 提取第一个问题
    first_question = None
    answers = []

    for entry in conversation_history:
        if entry.startswith('Q: ') and first_question is None:
            first_question = entry[3:]  # 去除 "Q: " 前缀
        elif entry.startswith('A: '):
            answers.append(entry[3:])  # 去除 "A: " 前缀

    if first_question:
        doc.add_paragraph(first_question)
        doc.add_paragraph()

    # 导出所有答案（只输出答案内容，不加任何标签）
    if answers:
        for answer in answers:
            # 只输出答案内容，不添加任何标签
            markdown_to_word(doc, answer)
            doc.add_paragraph()  # 添加空行分隔不同的答案

    # 如果没有找到内容
    if not first_question and not answers:
        doc.add_paragraph("No essential content found to export.")

    # 保存到内存
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    return file_stream


async def fetch_url_api(url: str, table_name: str = "uploaded_data",
                        use_javascript: bool = False,
                        wait_for_selector: Optional[str] = None,
                        timeout: int = 30):
    """调用后端接口爬取网页内容"""
    api_url = f"http://127.0.0.1:{config_data['server_port']}/fetch-url/"

    # 构建表单数据
    data = {
        "url": url,
        "table_name": table_name,
        "use_javascript": str(use_javascript).lower(),
        "timeout": str(timeout)
    }
    if wait_for_selector:
        data["wait_for_selector"] = wait_for_selector

    with httpx.Client(timeout=float(timeout) + 10) as client:
        try:
            response = client.post(api_url, data=data)
            if response.status_code == 200:
                return response.json()
            else:
                return {"error": f"HTTP {response.status_code}", "details": response.text}
        except httpx.RequestError as e:
            return {"error": f"Request failed: {str(e)}"}


def handle_url_fetch():
    """处理网页URL爬取"""
    # 获取URL输入
    url = input("🌐 Enter URL to fetch:", type=TEXT,
                placeholder="https://example.com", required=True)

    if not url:
        toast("❌ URL is required!", color='warning')
        return

    table_name = input(
        "📋 Table name (optional):",
        type=TEXT,
        placeholder="uploaded_data",
        required=False
    )
    if not table_name:
        table_name = "uploaded_data"

    # 使用默认参数：不启用JavaScript，超时30秒
    use_javascript = True
    wait_for_selector = None
    timeout = 60

    with put_loading(shape="grow", color="primary"):
        # 使用 asyncio.run() 来运行异步函数
        import asyncio
        result = asyncio.run(fetch_url_api(url, table_name, use_javascript, wait_for_selector, timeout))

    if result.get('error'):
        toast(f"❌ Fetch failed: {result.get('error')}", color='error')
        if result.get('details'):
            put_markdown(f"**Details:** {result.get('details')}")
    else:
        toast("✅ URL fetched successfully!", color='success')
        put_markdown("### 📊 Fetch Results")
        put_markdown(f"**URL:** {url}")
        put_markdown(f"**Table name:** `{result.get('table_name', table_name)}`")
        put_markdown(f"**Extracted text length:** {result.get('extracted_text_length', 'N/A')}")

        # 显示元数据
        metadata = result.get('metadata', {})
        if metadata:
            with put_collapse("📋 Page Metadata"):
                metadata_table = [["Property", "Value"]]
                for key, value in metadata.items():
                    if value:
                        metadata_table.append([key, str(value)[:200]])
                put_table(metadata_table)

        # 显示预览
        preview = result.get('preview', '')
        if preview:
            with put_collapse("📄 Content Preview"):
                put_markdown(preview)


def handle_export_word(conversation_history, export_type="full"):
    """处理导出Word文档 - 直接下载"""
    if not conversation_history:
        toast("⚠️ No content to export!", color='warning')
        return

    with put_loading(shape="grow", color="primary"):
        try:
            if export_type == "full":
                word_file = export_full_to_word(conversation_history)
                filename = "conversation_export_full.docx"
            else:  # essentials
                word_file = export_essentials_to_word(conversation_history)
                filename = "conversation_export_essentials.docx"

            # 读取文件内容并编码为base64用于直接下载
            file_content = word_file.getvalue()
            b64 = base64.b64encode(file_content).decode()

            # 使用JavaScript触发直接下载
            download_script = f'''
            <script>
                var link = document.createElement('a');
                link.href = 'data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}';
                link.download = '{filename}';
                document.body.appendChild(link);
                link.click();
                document.body.removeChild(link);
            </script>
            '''

            put_html(download_script)
            toast(f"✅ Download started: {filename}", color='success')

        except Exception as e:
            toast(f"❌ Export failed: {str(e)}", color='error')


def handle_csv_upload():
    file_info = file_upload(
        "📊 Please select a CSV file to upload",
        accept=".csv",
        help_text="Select the CSV file you want to upload"
    )

    if file_info:
        table_name = input("📋 Enter table name (optional, default is 'uploaded_data')", type=TEXT,
                           placeholder="uploaded_data", required=False)
        if not table_name:
            table_name = "uploaded_data"
        with put_loading(shape="grow", color="primary"):
            result = upload_csv_api(file_info['content'], table_name)
            print(result)
        err = result.get('type', "error")
        if err == "error":
            toast(f"❌ Upload failed: {result}", color='error')
        else:
            toast("✅ File uploaded successfully!", color='success')
            put_markdown("### 📊 Upload Results")
            put_markdown(f"Table name: `{result.get('table_name', table_name)}`")
            put_markdown(f"Row count: {result.get('row_count', 'N/A')}")
            put_markdown(f"Message: {result.get('message', 'N/A')}")


def handle_doc_upload():
    file_info = file_upload(
        "📄 Please select a document file to upload (txt, doc, docx, pdf)",
        accept=".txt,.doc,.docx,.pdf",
        help_text="Select the document file you want to upload"
    )

    if file_info:
        table_name = input("📋 Enter table name (optional, default is 'uploaded_data')", type=TEXT,
                           placeholder="uploaded_data", required=False)
        if not table_name:
            table_name = "uploaded_data"
        with put_loading(shape="grow", color="primary"):
            result = upload_doc_api(file_info['content'], file_info['filename'], table_name)
            print(result)

        if result.get('error'):
            toast(f"❌ Upload failed: {result.get('error')}", color='error')
        else:
            toast("✅ File uploaded successfully!", color='success')
            put_markdown("### 📄 Upload Results")
            put_markdown(f"Table name: `{result.get('table_name', table_name)}`")
            put_markdown(f"Extracted text length: {result.get('extracted_text_length', 'N/A')}")
            put_markdown(f"Preview: {result.get('preview', 'N/A')}")


def handle_table_selection(table_options):
    global SELECT_TABLES, SELECT_LABELS
    checkbox_options = [(opt['label'], opt['value']) for opt in table_options]
    selected_tables = checkbox(
        "📋 Select tables: ",
        options=checkbox_options,
        inline=True
    )
    SELECT_TABLES = selected_tables
    put_markdown(f"✅ You have selected: `{', '.join(selected_tables)}`")
    if selected_tables:
        selected_labels = []
        for table_value in selected_tables:
            for opt in table_options:
                if opt['value'] == table_value:
                    selected_labels.append(opt['label'])
                    break
        SELECT_LABELS = selected_labels


def create_feature_card(title, description, icon, button_text, onclick_func, scope_name):
    """创建功能卡片 - 使用table方式布局"""
    # 创建scope用于放置卡片
    put_scope(scope_name)

    with use_scope(scope_name):
        # 卡片内容
        card_html = f'''
        <div style="
            border: 1px solid #ddd;
            border-radius: 8px;
            padding: 15px;
            margin: 5px;
            text-align: center;
            background: #f8f9fa;
            box-shadow: 0 2px 4px rgba(0,0,0,0.1);
            min-height: 180px;
            display: flex;
            flex-direction: column;
            justify-content: center;
        ">
            <div style="font-size: 48px; margin-bottom: 10px;">{icon}</div>
            <h3 style="margin: 10px 0; font-size: 16px;">{title}</h3>
            <p style="color: #666; font-size: 13px; margin: 5px 0;">{description}</p>
        </div>
        '''
        put_html(card_html)
        # 按钮放在卡片下方
        put_buttons([button_text], onclick=[onclick_func])


def main():
    global SELECT_TABLES, SELECT_LABELS
    conversation_history = []

    # 设置页面标题
    put_markdown("# 🤖 Medical Dataset Copilot")
    put_markdown("---")

    # 获取表格信息
    table_comments = get_table_comments_dict()
    table_options = []
    for table_name, comment in table_comments.items():
        display_name = f"{table_name} ({comment})" if comment else table_name
        table_options.append({'label': display_name, 'value': table_name})

    put_table([
        [
            # 第一列：CSV上传
            put_scope('csv_card'),
            # 第二列：文档上传
            put_scope('doc_card'),
            # 第三列：URL抓取
            put_scope('url_card')
        ]
    ])

    # 读取图片为base64格式
    def get_image_base64(image_path):
        try:
            with open(image_path, 'rb') as f:
                img_data = f.read()
                return base64.b64encode(img_data).decode()
        except Exception as e:
            print(f"Error loading image {image_path}: {e}")
            return None

    # 获取三张图片的base64
    csv_img_base64 = get_image_base64('./front_img/upload_csv.png')
    doc_img_base64 = get_image_base64('./front_img/upload_doc.png')
    url_img_base64 = get_image_base64('./front_img/fetch_url.png')

    # 填充三个卡片的内容
    # 第一列：CSV上传
    with use_scope('csv_card'):
        put_html('''
        <div style="
            border: 1px solid #ddd;
            border-radius: 8px;
            padding: 15px;
            margin: 5px;
            text-align: center;
            background: #f8f9fa;
            box-shadow: 0 2px 4px rgba(0,0,0,0.1);
            min-height: 160px;
            display: flex;
            flex-direction: column;
            justify-content: space-between;
            align-items: center;
        ">
            <div style="width: 100%;">
                <div style="font-size: 48px; margin-bottom: 10px;">📊</div>
                <h3 style="margin: 10px 0; font-size: 16px;">CSV Upload</h3>
                <p style="color: #666; font-size: 13px; margin: 5px 0;">Upload CSV file for data analysis</p>
            </div>
        </div>
        ''')
        # 在卡片下方显示图片
        if csv_img_base64:
            put_html(
                f'<div style="height:160px ;text-align: center; margin: 5px 0;"><img src="data:image/png;base64,{csv_img_base64}" style="max-width: 240px; max-height: 240px;"/></div>')
        put_buttons([' 📤 Upload CSV'], onclick=[handle_csv_upload])

    # 第二列：文档上传
    with use_scope('doc_card'):
        put_html('''
        <div style="
            border: 1px solid #ddd;
            border-radius: 8px;
            padding: 15px;
            margin: 5px;
            text-align: center;
            background: #f8f9fa;
            box-shadow: 0 2px 4px rgba(0,0,0,0.1);
            min-height: 160px;
            display: flex;
            flex-direction: column;
            justify-content: space-between;
            align-items: center;
        ">
            <div style="width: 100%;">
                <div style="font-size: 48px; margin-bottom: 10px;">📄</div>
                <h3 style="margin: 10px 0; font-size: 16px;">Document Upload</h3>
                <p style="color: #666; font-size: 13px; margin: 5px 0;">Upload TXT, DOC, DOCX, or PDF</p>
            </div>
        </div>
        ''')
        # 在卡片下方显示图片
        if doc_img_base64:
            put_html(
                f'<div style="height:160px ;text-align: center; margin: 5px 0;"><img src="data:image/png;base64,{doc_img_base64}" style="max-width: 240px; max-height: 240px;"/></div>')
        put_buttons(['📤 Upload Document'], onclick=[handle_doc_upload])

    # 第三列：URL抓取
    with use_scope('url_card'):
        put_html('''
        <div style="
            border: 1px solid #ddd;
            border-radius: 8px;
            padding: 15px;
            margin: 5px;
            text-align: center;
            background: #f8f9fa;
            box-shadow: 0 2px 4px rgba(0,0,0,0.1);
            min-height: 160px;
            display: flex;
            flex-direction: column;
            justify-content: space-between;
            align-items: center;
        ">
            <div style="width: 100%;">
                <div style="font-size: 48px; margin-bottom: 10px;">🌐</div>
                <h3 style="margin: 10px 0; font-size: 16px;">Fetch URL</h3>
                <p style="color: #666; font-size: 13px; margin: 5px 0;">Extract content from web pages</p>
            </div>
        </div>
        ''')
        # 在卡片下方显示图片
        if url_img_base64:
            put_html(
                f'<div style="height:160px ;text-align: center; margin: 5px 0;"><img src="data:image/png;base64,{url_img_base64}" style="max-width: 240px; max-height: 240px;"/></div>')
        put_buttons([' 🌐 Fetch URL  '], onclick=[handle_url_fetch])

    put_markdown("---")


    put_markdown("### 📥 Export Options")
    put_buttons(['Export Full Conversation', '📝 Export Essentials (Answers)'],
                onclick=[lambda: handle_export_word(conversation_history, "full"),
                         lambda: handle_export_word(conversation_history, "essentials")])

    put_markdown("### 📊 Data View")
    with put_collapse(f"📋 Tables"):
        # 获取所有注释信息
        all_comments = get_all_comments_from_table()
        first_five_rows = get_rows_from_all_tables()

        for table_name, rows in first_five_rows.items():
            with put_collapse(f" table {table_name}"):
                # 显示表注释
                if table_name in all_comments:
                    table_comment = all_comments[table_name].get('table_comment', '')
                    if table_comment:
                        put_text(f"📝 {table_comment}")

                    # 显示列注释（表格形式）
                    columns = all_comments[table_name].get('columns', {})
                    if columns:
                        comment_table = [["Column Name", "Comment"]]
                        for col_name, comment in columns.items():
                            comment_table.append([col_name, comment])
                        put_table(comment_table)

                put_text(f"📊 table {table_name} first 5 rows:")
                put_table([rows.columns.tolist()] + rows.values.tolist())

    # 对话主循环
    while True:
        question = textarea("💬 Enter your question here:", type=TEXT, rows=2)
        put_markdown("## ❓ " + question)
        conversation_history.append(f"Q: {question}")
        with put_loading():
            step_str, _ = ai_agent_api(question, SELECT_TABLES, "/api/step-chat/")
        if step_str:
            # step_str = textarea("revise plan:", type=TEXT, rows=8, value=step_str)
            conversation_history.append(f"Planner: {step_str}")
            put_markdown(step_str, sanitize=False)
        else:
            put_text("❌ Failed to get a response from the AI Agent.")

        while True:
            table_pre = ""

            value = "please do the next step on the todo list"
            question = textarea("⏭ What is next?:", value=value, type=TEXT, rows=2)
            # question = value
            put_markdown("## ⏭️ " + question)
            if conversation_history:
                context = "\n".join(conversation_history)
                full_question = f"Context:\n{context}\n\nCurrent Question:\n{question}"
            else:
                full_question = question

            if value == question:
                with put_loading():
                    response, code = ai_agent_api(table_pre + full_question, SELECT_TABLES, "/api/ask-agent/")
                if response:
                    conversation_history.append(f"Q: {question}")
                    conversation_history.append(f"Code Generated: {code}")
                    conversation_history.append(f"A: {response}")
                    put_markdown(response, sanitize=False)
                    time.sleep(3)
                else:
                    put_text("❌ Failed to get a response from the AI Agent.")

                context = "\n".join(conversation_history)
                full_question = f"Context:\n{context}\n"
            else:
                context = "\n".join(conversation_history)
                full_question = f"Context:\n{context}\n\nCurrent Question:\n{question}"

            with put_loading():
                step_str, _ = ai_agent_api(table_pre + full_question, SELECT_TABLES, "/api/step-chat/")
            if step_str:
                # step_str = textarea("revise plan:", type=TEXT, rows=8, value=step_str)
                conversation_history.append(f"Planner: {step_str}")
                put_markdown(step_str, sanitize=False)
                if "[ ]" not in step_str:
                    break
            else:
                put_text("❌ Failed to get a response from the AI Agent.")


if __name__ == '__main__':
    start_server(main, port=8038, debug=True)
